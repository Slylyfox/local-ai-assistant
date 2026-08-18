"""Formats findings into a saved report file.

Deliberately stateless: the caller (the model, which already has the full
conversation in context) supplies the finished content rather than this tool
reading a session from the database. An ambient "current session" lookup
would race — the GUI and Telegram bridge can each have a different active
session generating at the same time — so keeping this tool a pure formatter
avoids that hazard entirely."""

import os
import re
import time

from config import DATA_DIR

REPORTS_DIR = os.path.join(DATA_DIR, "reports")


def _slugify(title: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title.strip()).strip("-").lower()
    return slug or "report"


def generate_report(title: str, content_markdown: str, format: str = "markdown") -> str:
    """Save formatted findings as a titled report file."""
    os.makedirs(REPORTS_DIR, exist_ok=True)
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    slug = _slugify(title)
    filename_base = f"{slug}_{time.strftime('%Y%m%d_%H%M%S')}"

    header = f"# {title}\n\n_Generated {timestamp}_\n\n---\n\n"
    full_markdown = header + content_markdown

    saved_paths = []

    if format in ("markdown", "both"):
        md_path = os.path.join(REPORTS_DIR, f"{filename_base}.md")
        try:
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(full_markdown)
            saved_paths.append(md_path)
        except OSError as exc:
            return f"Failed to write markdown report: {exc}"

    if format in ("docx", "both"):
        try:
            from docx import Document
        except ImportError:
            return "docx format requires the 'python-docx' package (pip install python-docx)." + (
                f" Markdown report was still saved to '{saved_paths[0]}'." if saved_paths else ""
            )
        try:
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"Generated {timestamp}").italic = True
            for line in content_markdown.split("\n"):
                stripped = line.strip()
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith(("- ", "* ")):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    doc.add_paragraph(stripped)
            docx_path = os.path.join(REPORTS_DIR, f"{filename_base}.docx")
            doc.save(docx_path)
            saved_paths.append(docx_path)
        except Exception as exc:  # noqa: BLE001
            return f"Failed to write docx report: {exc}" + (
                f" Markdown report was still saved to '{saved_paths[0]}'." if saved_paths else ""
            )

    if not saved_paths:
        return f"Unknown format '{format}'. Use 'markdown', 'docx', or 'both'."
    return "Report saved to: " + ", ".join(f"'{p}'" for p in saved_paths)


def register(registry):
    registry.register(
        "generate_report",
        "Format already-synthesized findings into a titled report and save it locally as markdown "
        "and/or Word (.docx). Pass the full report body as content_markdown — this tool only formats "
        "and saves, it doesn't gather or summarize content itself.",
        {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Report title."},
                "content_markdown": {
                    "type": "string",
                    "description": "The report body, in markdown (headers, bullets, paragraphs).",
                },
                "format": {
                    "type": "string",
                    "description": "One of 'markdown', 'docx', or 'both'. Defaults to 'markdown'.",
                },
            },
            "required": ["title", "content_markdown"],
        },
        generate_report,
        category="report",
    )
